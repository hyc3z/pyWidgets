from docx import Document
from docx2python import docx2python
import os

#读取文档
def readTable(filename, index_cell = (3,)):
    doc = Document(filename) #filename为word文档
    tables = doc.tables  #获取文档的表格个数 len(doc.tables)
    header_cells = tables[0].rows[0].cells
    headers = [x.text for x in header_cells]
    header_text = [header_cells[x].text for x in index_cell]
    header_str = "-".join(header_text)
    data = {}
    # print(headers)
    for table in tables:
        rows = table.rows  #获取表格的行数len(tb1.rows)
        for row in rows:
            row_cells = row.cells
            cell_text = [row_cells[x].text for x in index_cell]
            index_str = "-".join(cell_text)
            curdict = {}
            for i in range(len(headers)):
                cell = row_cells[i]
                curdict[headers[i]] = cell.text
            if index_str != header_str:
                data[index_str] = curdict
    return data


def getPicture(filename, storedir):
    doc = docx2python(filename)
    for name, image in doc.images.items():
        with open(os.path.join(storedir, name), 'wb') as fp:
            fp.write(image)

def main():
    # getPicture("src.docx", "pic")
    table_dict = readTable("src.docx")

if __name__ == '__main__':
    main()